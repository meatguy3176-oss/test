"""
사고조사보고서 .docx 생성 모듈
- 2페이지 구성: 1페이지 본문 / 2페이지 위치도+현장사진
- 이미지는 박스 안에 비율 유지하며 자동 축소
"""
import os
from PIL import Image as PILImage
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_korean_font(run, font_name='맑은 고딕', size=11, bold=True, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color_rgb is not None:
        run.font.color.rgb = color_rgb
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_text_paragraph(doc, text, size=11, bold=True, align='left', color_rgb=None,
                       space_before=0, space_after=0, line_spacing=None, underline=False):
    p = doc.add_paragraph()
    align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                 'center': WD_ALIGN_PARAGRAPH.CENTER,
                 'right': WD_ALIGN_PARAGRAPH.RIGHT}
    p.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    if underline:
        run.underline = True
    set_korean_font(run, size=size, bold=bold, color_rgb=color_rgb)
    return p


def fit_picture_size(img_path, max_width_cm, max_height_cm):
    """이미지 비율 유지하며 박스 안에 맞는 크기(Cm) 계산."""
    with PILImage.open(img_path) as img:
        w_px, h_px = img.size
        dpi = img.info.get('dpi', (96, 96))
        dpi_x = dpi[0] if dpi[0] > 0 else 96
        dpi_y = dpi[1] if dpi[1] > 0 else 96
    w_cm = w_px / dpi_x * 2.54
    h_cm = h_px / dpi_y * 2.54
    scale = min(max_width_cm / w_cm, max_height_cm / h_cm)
    return Cm(w_cm * scale), Cm(h_cm * scale)


def add_fitted_picture(paragraph, img_path, max_width_cm, max_height_cm):
    w, h = fit_picture_size(img_path, max_width_cm, max_height_cm)
    run = paragraph.add_run()
    run.add_picture(img_path, width=w, height=h)


def create_accident_report(data, output_path):
    """
    사고조사보고서 docx 생성.
    data: dict (accident_datetime, accident_place, complex_no, damage_person,
                damage_property, responders, actions[list], cause, measures,
                location_img[path|None], site_imgs[list of paths])
    output_path: 저장할 .docx 절대경로
    """
    doc = Document()

    # ----- 페이지 설정 (A4) -----
    for section in doc.sections:
        section.top_margin = Mm(18)
        section.bottom_margin = Mm(18)
        section.left_margin = Mm(22)
        section.right_margin = Mm(22)

    # ========== 1페이지: 본문 ==========
    add_text_paragraph(doc, '사고 조사 보고서', size=20, bold=True, align='center',
                       space_after=10)

    # 상단 헤더 표 (얇은 회색 밴드)
    header_table = doc.add_table(rows=1, cols=1)
    header_table.style = 'Table Grid'
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.rows[0].cells[0]
    header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    tcPr = header_cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9D9D9')
    tcPr.append(shd)

    header_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    header_table.rows[0].height = Pt(4)
    header_p = header_cell.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.space_before = Pt(0)
    header_p.paragraph_format.space_after = Pt(0)
    header_p.paragraph_format.line_spacing = Pt(1)
    run = header_p.add_run(' ')
    set_korean_font(run, size=1)

    SECTION_GAP = 25
    LINE_SP = 1

    # 1. 사고일시
    add_text_paragraph(doc, f"1. 사고일시 : {data['accident_datetime']}", size=11,
                       space_before=SECTION_GAP, line_spacing=LINE_SP)

    # 2. 사고장소
    add_text_paragraph(doc, f"2. 사고장소 : {data['accident_place']}", size=11,
                       space_before=SECTION_GAP, line_spacing=LINE_SP)
    add_text_paragraph(doc, f"             (동단지번호 : {data['complex_no']})", size=11,
                       line_spacing=LINE_SP)

    # 3. 피해현황
    add_text_paragraph(doc, "3. 피해현황", size=11,
                       space_before=SECTION_GAP, line_spacing=LINE_SP)
    add_text_paragraph(doc, f"    - 인명 : {data['damage_person']}", size=11, line_spacing=LINE_SP)
    add_text_paragraph(doc, f"    - 재산 : {data['damage_property']}", size=11, line_spacing=LINE_SP)

    # 4. 사고내용(시간대별 조치사항)
    add_text_paragraph(doc, "4. 사고내용(시간대별 조치사항)", size=11,
                       space_before=SECTION_GAP, line_spacing=LINE_SP)
    actions = data.get('actions', [])
    if actions:
        first = actions[0]
        add_text_paragraph(doc, f"   ◦ {first['time']} : {first['content']}", size=11,
                           line_spacing=LINE_SP)
        if data.get('responders'):
            add_text_paragraph(doc, f"             (출동자 : {data['responders']})", size=11,
                               line_spacing=LINE_SP)
        for action in actions[1:]:
            add_text_paragraph(doc, f"   ◦ {action['time']} : {action['content']}", size=11,
                               line_spacing=LINE_SP)

    # 5. 사고원인
    add_text_paragraph(doc, f"5. 사고원인 : {data['cause']}", size=11,
                       space_before=SECTION_GAP, line_spacing=LINE_SP)

    # 6. 조치사항
    add_text_paragraph(doc, "6. 조치사항", size=11,
                       space_before=SECTION_GAP, line_spacing=LINE_SP)
    measures_list = [m.strip() for m in data['measures'].split('\n') if m.strip()]
    for measure in measures_list:
        add_text_paragraph(doc, f"    - {measure}", size=11, line_spacing=LINE_SP)

    # 붙임
    add_text_paragraph(doc, "붙임", size=11,
                       space_before=SECTION_GAP, line_spacing=LINE_SP)
    add_text_paragraph(doc, " 1. 위치도", size=11, line_spacing=LINE_SP)
    add_text_paragraph(doc, " 2. 현장사진", size=11, line_spacing=LINE_SP)

    # ========== 2페이지: 위치도 + 현장사진 ==========
    doc.add_page_break()

    add_text_paragraph(doc, "위   치   도", size=14, bold=True, align='center',
                       space_after=8, underline=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if data.get('location_img') and os.path.exists(data['location_img']):
        add_fitted_picture(p, data['location_img'],
                           max_width_cm=16.0, max_height_cm=10.0)
    else:
        run = p.add_run("[위치도 이미지 없음]")
        set_korean_font(run, size=11)

    add_text_paragraph(doc, "현 장 사 진", size=14, bold=True, align='center',
                       space_before=14, space_after=8, underline=True)

    site_imgs = data.get('site_imgs', [])
    site_table = doc.add_table(rows=2, cols=2)
    site_table.style = 'Table Grid'
    site_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    CELL_MAX_W = 7.5
    CELL_MAX_H = 5.0

    for i in range(4):
        row_idx, col_idx = divmod(i, 2)
        cell = site_table.rows[row_idx].cells[col_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.width = Cm(8.0)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i < len(site_imgs) and site_imgs[i] and os.path.exists(site_imgs[i]):
            add_fitted_picture(p, site_imgs[i],
                               max_width_cm=CELL_MAX_W, max_height_cm=CELL_MAX_H)
        else:
            run = p.add_run(f"[사진 {i+1} 없음]")
            set_korean_font(run, size=10)

    doc.save(output_path)
    return output_path
